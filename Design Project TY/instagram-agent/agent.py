import os
import time
import re
from datetime import datetime
import streamlit as st
import pandas as pd
from pytrends.request import TrendReq
from googleapiclient.discovery import build
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

load_dotenv()

GROQ_API_KEY    = os.getenv("GROQ_API_KEY")
YOUTUBE_API_KEY = os.getenv("YOUTUBE_API_KEY")

# ─── GROQ CLIENT ──────────────────────────────────────────────
def get_groq_client():
    if not GROQ_API_KEY:
        st.error("❌ GROQ_API_KEY missing in .env file")
        return None
    return Groq(api_key=GROQ_API_KEY)


# ─── GOOGLE TRENDS ────────────────────────────────────────────
def fetch_google_trends(keywords):
    results = []
    try:
        pytrends = TrendReq(hl='en-US', tz=330)
        for kw in keywords:
            try:
                pytrends.build_payload([kw], timeframe='now 7-d', geo='IN')
                data = pytrends.interest_over_time()
                if not data.empty:
                    score = float(data[kw].mean())
                    results.append({
                        "keyword": kw,
                        "score": score,
                        "source": "Google Trends"
                    })
                time.sleep(1.5)
            except Exception:
                continue
    except Exception as e:
        st.warning(f"Google Trends issue: {e}")
    return results


# ─── YOUTUBE TRENDING ─────────────────────────────────────────
def fetch_youtube_trends(max_results=10):
    results = []
    if not YOUTUBE_API_KEY:
        st.warning("⚠️ No YouTube API key — skipping YouTube trends")
        return results
    try:
        youtube = build("youtube", "v3", developerKey=YOUTUBE_API_KEY)
        response = youtube.videos().list(
            part="snippet,statistics",
            chart="mostPopular",
            regionCode="IN",
            maxResults=max_results
        ).execute()

        for item in response.get("items", []):
            title = item["snippet"]["title"]
            views = int(item.get("statistics", {}).get("viewCount", 0))
            results.append({
                "keyword": title[:60],
                "score": views / 100000,
                "source": "YouTube"
            })
    except Exception as e:
        st.warning(f"YouTube issue: {e}")
    return results


# ─── MERGE + RANK TRENDS ──────────────────────────────────────
def get_top_trends(google_results, youtube_results, top_n=3):
    all_trends = google_results + youtube_results
    if not all_trends:
        return []
    df = pd.DataFrame(all_trends)
    df = df.sort_values("score", ascending=False).drop_duplicates("keyword")
    return df.head(top_n).to_dict("records")


# ─── GROQ AI GENERATION ───────────────────────────────────────
def generate_post_ideas(client, trends, niche="Tech / AI / Education"):
    trend_list = "\n".join([
        f"- {t['keyword']} (score: {t['score']:.1f}, source: {t['source']})"
        for t in trends
    ])

    prompt = f"""
You are an expert Instagram content strategist specializing in {niche}.

Here are this week's top trending topics:
{trend_list}

For EACH trend, create a complete Instagram post package.

Format your response EXACTLY like this for each post:

---POST 1---
TREND: [trend name]
POST TYPE: [Carousel / Reel / Single Image]
TITLE: [catchy title for the post]
HOOK: [first line that stops the scroll — max 10 words]
CAPTION: [full caption, 3-5 sentences, conversational and engaging]
SLIDE IDEAS: [if carousel: list 5-6 slide titles. if reel: list 3-4 scene ideas]
CALL TO ACTION: [one clear CTA]
HASHTAGS: [15 hashtags, mix of popular and niche, with #]
BEST TIME TO POST: [day and time in IST]
PRO TIP: [one insider tip to boost this post's reach]

---POST 2---
[same format]

---POST 3---
[same format]

Make each post unique in format (one carousel, one reel, one single image).
Be very specific, creative and actionable. Write like a real content creator.
"""

    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.8,
            max_tokens=3000
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Groq generation error: {e}")
        return None


# ─── PARSE AI OUTPUT ──────────────────────────────────────────
def parse_posts(raw_text):
    posts = []
    blocks = re.split(r'---POST \d+---', raw_text)
    blocks = [b.strip() for b in blocks if b.strip()]

    for block in blocks:
        post = {}
        fields = [
            "TREND", "POST TYPE", "TITLE", "HOOK", "CAPTION",
            "SLIDE IDEAS", "CALL TO ACTION", "HASHTAGS",
            "BEST TIME TO POST", "PRO TIP"
        ]
        for i, field in enumerate(fields):
            if i < len(fields) - 1:
                pattern = field + r':\s*(.*?)(?=' + '|'.join(fields[i+1:]) + r'|$)'
            else:
                pattern = field + r':\s*(.*?)$'
            match = re.search(pattern, block, re.DOTALL | re.IGNORECASE)
            if match:
                post[field] = match.group(1).strip()
            else:
                post[field] = ""
        if post.get("TREND") or post.get("TITLE"):
            posts.append(post)
    return posts


# ─── CREATE WORD DOC ──────────────────────────────────────────
def create_word_doc(posts, trends, week_date):
    doc = Document()

    section = doc.sections[0]
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

    # Title
    title = doc.add_heading('', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Weekly Instagram Content Plan')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        f'Week of {week_date}  |  Niche: Tech / AI / Education'
    )
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    sub_run.font.italic = True

    doc.add_paragraph()

    # Trending Topics Summary
    doc.add_heading("This Week's Trending Topics", level=1)
    for i, t in enumerate(trends, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"  {i}. {t['keyword']}")
        r.bold = True
        p.add_run(f"  —  Score: {t['score']:.1f}  |  Source: {t['source']}")

    doc.add_paragraph()
    doc.add_paragraph('=' * 60)
    doc.add_paragraph()

    # Posts
    for i, post in enumerate(posts, 1):
        h = doc.add_heading('', level=1)
        h_run = h.add_run(f"POST {i}  —  {post.get('POST TYPE', '').upper()}")
        h_run.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)

        def add_field(label, key, bold_value=False):
            value = post.get(key, "")
            if value:
                p = doc.add_paragraph()
                label_run = p.add_run(f"{label}  ")
                label_run.bold = True
                label_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
                val_run = p.add_run(value)
                val_run.bold = bold_value
                val_run.font.size = Pt(11)

        add_field("Trend:",          "TREND",           bold_value=True)
        add_field("Title:",          "TITLE",           bold_value=True)
        add_field("Hook:",           "HOOK",            bold_value=True)
        doc.add_paragraph()
        add_field("Caption:",        "CAPTION")
        doc.add_paragraph()
        add_field("Content Ideas:",  "SLIDE IDEAS")
        doc.add_paragraph()
        add_field("Call to Action:", "CALL TO ACTION")
        add_field("Hashtags:",       "HASHTAGS")
        add_field("Best Time:",      "BEST TIME TO POST")
        add_field("Pro Tip:",        "PRO TIP")

        doc.add_paragraph()
        doc.add_paragraph('=' * 60)
        doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer.add_run(
        f'Generated by AI Content Agent  |  {datetime.now().strftime("%d %b %Y, %I:%M %p")}'
    )
    f_run.font.size = Pt(9)
    f_run.font.italic = True
    f_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    filename = f"Weekly_Posts_{datetime.now().strftime('%d_%m_%Y')}.docx"
    filepath = os.path.join("outputs", filename)
    os.makedirs("outputs", exist_ok=True)
    doc.save(filepath)
    return filepath, filename


# ─── STREAMLIT UI ─────────────────────────────────────────────
def main():
    st.set_page_config(
        page_title="Instagram Content Agent",
        page_icon="📱",
        layout="centered"
    )

    st.title("📱 Instagram Content Agent")
    st.markdown(
        "*Scrapes trending topics → Generates 3 ready-to-use post ideas → Downloads as Word doc*"
    )
    st.divider()

    # Sidebar
    with st.sidebar:
        st.header("⚙️ Settings")
        niche = st.selectbox(
            "Your Instagram Niche",
            [
                "Tech / AI / Education",
                "Finance / Business",
                "Fitness / Health",
                "Entertainment / Memes"
            ]
        )
        st.markdown("---")
        st.markdown("**Keywords to Track**")
        custom_input = st.text_area(
            "One keyword per line",
            value="AI\nmachine learning\nchatgpt\npython programming\ntech news",
            height=150
        )
        st.markdown("---")
        st.info("Runs take 30–60 seconds due to Google Trends rate limits.")

    keywords = [k.strip() for k in custom_input.strip().split("\n") if k.strip()]

    # Main button
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        run_btn = st.button(
            "🚀 Generate This Week's Posts",
            use_container_width=True,
            type="primary"
        )

    if run_btn:
        st.divider()

        # Step 1 — collect trends
        with st.status("🔍 Collecting trending data...", expanded=True) as status:
            st.write("📈 Fetching Google Trends...")
            google_data = fetch_google_trends(keywords)
            st.write(f"✅ Google Trends: {len(google_data)} keywords collected")

            st.write("🎥 Fetching YouTube trending...")
            youtube_data = fetch_youtube_trends()
            st.write(f"✅ YouTube: {len(youtube_data)} videos collected")

            top_trends = get_top_trends(google_data, youtube_data, top_n=3)

            if not top_trends:
                st.error("❌ No trends found. Check your API keys.")
                st.stop()

            st.write("🏆 Top 3 trends identified!")
            status.update(label="✅ Trend data collected!", state="complete")

        # Show trends
        st.subheader("🔥 This Week's Top Trends")
        cols = st.columns(3)
        for i, trend in enumerate(top_trends):
            with cols[i]:
                st.metric(
                    label=f"#{i+1} — {trend['source']}",
                    value=trend['keyword'][:30],
                    delta=f"Score: {trend['score']:.1f}"
                )

        st.divider()

        # Step 2 — AI generation
        with st.status("🤖 Generating post ideas with Groq AI...", expanded=True) as status:
            client = get_groq_client()
            if not client:
                st.stop()

            st.write("✍️ Creating 3 complete post packages...")
            raw_output = generate_post_ideas(client, top_trends, niche)

            if not raw_output:
                st.error("❌ AI generation failed. Check your Groq API key.")
                st.stop()

            posts = parse_posts(raw_output)
            st.write(f"✅ {len(posts)} post ideas generated!")
            status.update(label="✅ Posts generated!", state="complete")

        # Preview posts
        st.subheader("📋 Post Ideas Preview")
        for i, post in enumerate(posts, 1):
            with st.expander(
                f"Post {i} — {post.get('POST TYPE', '')} | {post.get('TITLE', 'View Post')}",
                expanded=(i == 1)
            ):
                col_a, col_b = st.columns(2)
                with col_a:
                    st.markdown("**🪝 Hook**")
                    st.info(post.get("HOOK", ""))
                    st.markdown("**✍️ Caption**")
                    st.write(post.get("CAPTION", ""))
                with col_b:
                    st.markdown("**🎬 Content Ideas**")
                    st.write(post.get("SLIDE IDEAS", ""))
                    st.markdown("**🏷️ Hashtags**")
                    st.code(post.get("HASHTAGS", ""), language=None)
                st.markdown(
                    f"⏰ **Best Time:** {post.get('BEST TIME TO POST', '')}  "
                    f"|  💡 **Pro Tip:** {post.get('PRO TIP', '')}"
                )

        st.divider()

        # Step 3 — create doc
        with st.status("📄 Creating Word document...", expanded=True) as status:
            week_date = datetime.now().strftime("%d %B %Y")
            filepath, filename = create_word_doc(posts, top_trends, week_date)
            status.update(label="✅ Document ready!", state="complete")

        # Download
        st.success("🎉 Your weekly content plan is ready!")
        with open(filepath, "rb") as f:
            st.download_button(
                label="📥 Download Word Doc",
                data=f,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                type="primary"
            )

        st.balloons()


if __name__ == "__main__":
    main()